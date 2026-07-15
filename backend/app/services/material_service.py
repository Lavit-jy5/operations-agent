from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import UploadFile


SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx"}
MAX_FILE_BYTES = 3 * 1024 * 1024


async def extract_materials(files: list[UploadFile]) -> tuple[str, list[str]]:
    sections: list[str] = []
    issues: list[str] = []

    for file in files:
        filename = file.filename or "未命名文件"
        suffix = Path(filename).suffix.lower()
        content = await file.read()

        if suffix not in SUPPORTED_EXTENSIONS:
            issues.append(f"{filename}：暂不支持该格式，当前支持 txt、md、docx。")
            continue

        if len(content) > MAX_FILE_BYTES:
            issues.append(f"{filename}：文件超过 3MB，已跳过。")
            continue

        try:
            if suffix in {".txt", ".md"}:
                text = _decode_text(content)
            else:
                text = _extract_docx(content)
        except Exception as exc:
            issues.append(f"{filename}：解析失败，原因：{exc}")
            continue

        clean_text = _clean_text(text)
        if not clean_text:
            issues.append(f"{filename}：未提取到有效文字。")
            continue

        sections.append(f"【文件：{filename}】\n{clean_text}")

    return "\n\n".join(sections), issues


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _extract_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _clean_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    compact = [line for line in lines if line]
    return "\n".join(compact)[:12000]

import json
import re
from io import BytesIO
from typing import Optional
from docx import Document
from docx.shared import Pt

from app.core.config import settings
from app.models.schemas import (
    GeneratedContentRequest,
    QualityCheckResponse,
    TitleCandidatesRequest,
    TitleCandidatesResponse,
    ValidationIssue,
)
from app.services.bailian_client import BailianClientError, chat_completion


FORBIDDEN_SOURCE_WORDS = (
    "数据来源：",
    "来源：",
    "据材料",
    "材料显示",
    "材料正文",
    "文字内容",
    "图片材料",
    "图片解析材料",
    "图片识别内容",
    "文档解析材料",
    "参考资料内容",
)
FORBIDDEN_TITLE_WORDS = ("必然", "确定反转", "现在就买", "布局窗口", "收益")
DISALLOWED_COMPANIES = (
    "广发",
    "易方达",
    "华夏",
    "嘉实",
    "招商",
    "博时",
    "富国",
    "南方",
    "鹏华",
    "汇添富",
    "银华",
    "华安",
    "华宝",
    "万家",
    "景顺",
    "中欧",
)
PRODUCT_WORD_RE = re.compile(r"ETF|场内|场外|联接|LOF|QDII|指数基金|基金代码|A类|C类", re.I)


def quality_check(content: GeneratedContentRequest) -> QualityCheckResponse:
    issues = build_quality_report(content)
    score = _score_from_issues(issues)
    return QualityCheckResponse(quality_score=score, review_issues=issues)


def build_quality_report(content: GeneratedContentRequest) -> list[ValidationIssue]:
    text = "\n".join([content.title, content.summary, content.body, content.risk_notice, "\n".join(content.citations)])
    findings = trace_content_data(content)
    body_paragraphs = _split_paragraphs(content.body)
    issues: list[ValidationIssue] = []

    title = content.title.strip()
    title_ok = bool(title) and len(title) >= 8 and "草稿" not in title
    issues.append(
        _rule_issue(
            1,
            "标题完整性",
            "检查是否有标题，标题是否为空、过短或明显像草稿。",
            "通过" if title_ok else "未通过",
            f"标题已填写，长度为 {len(title)} 个字符。" if title_ok else "标题为空、过短，或仍包含草稿化表达。",
            f"标题：{title or '未填写'}",
        )
    )

    body_ok = len(content.body.strip()) >= 80 and len(body_paragraphs) >= 2
    issues.append(
        _rule_issue(
            2,
            "正文完整性",
            "检查正文是否为空，是否只有一句话，是否明显没有生成完整内容。",
            "通过" if body_ok else "未通过",
            f"正文共 {len(body_paragraphs)} 个段落，长度约 {len(content.body.strip())} 个字符。" if body_ok else "正文过短或段落不足，可能不是完整成稿。",
            f"正文段落数：{len(body_paragraphs)}",
        )
    )

    structure_status, structure_detail, structure_location = _check_structure(content)
    issues.append(
        _rule_issue(
            3,
            "核心结构完整性",
            "热点文章检查是否包含【热点解读】、【后市展望】、【相关产品】、【风险提示】；热点异动检查是否包含主推方案、机会解读、自查清单等关键结构。",
            structure_status,
            structure_detail,
            structure_location,
        )
    )

    markdown_hits = _find_markdown_hits(text)
    issues.append(
        _rule_issue(
            4,
            "Markdown / 非发布格式",
            "检查是否出现 #、过多 **、代码块符号等不适合直接发布的格式残留。",
            "通过" if not markdown_hits else "提醒",
            "未发现 Markdown 标题符号、代码块符号或明显非发布格式。" if not markdown_hits else "发现疑似 Markdown 或非发布格式残留。",
            "\n".join(markdown_hits) if markdown_hits else "无",
        )
    )

    backend_hits = [word for word in FORBIDDEN_SOURCE_WORDS if word in text]
    issues.append(
        _rule_issue(
            5,
            "后台字段残留",
            "检查是否出现“材料显示”“据材料”“图片识别内容”“文档解析材料”“参考资料内容”“文字内容”等后台处理词。",
            "通过" if not backend_hits else "提醒",
            "未发现后台字段残留。" if not backend_hits else "发现后台字段残留，需要改成自然来源表达。",
            "无" if not backend_hits else "、".join(backend_hits),
        )
    )

    filename_hits = _find_filename_hits(text)
    issues.append(
        _rule_issue(
            6,
            "文件名残留",
            "检查是否出现 .jpg、.png、.docx、WechatIMG、截屏2026 等上传文件名。",
            "通过" if not filename_hits else "提醒",
            "未发现图片或文档文件名残留。" if not filename_hits else "发现上传文件名残留。",
            "\n".join(filename_hits) if filename_hits else "无",
        )
    )

    data_points = [finding["data"] for finding in findings]
    provided_count = sum(1 for finding in findings if finding["status"] != "未找到")
    missing_count = len(findings) - provided_count
    issues.append(
        _rule_issue(
            7,
            "数据存在性",
            "检查正文中是否出现关键数据，并判断这些数据是否来自用户提供的文字材料、图片识别内容、参考文档或 Wind 数据。",
            _data_presence_status(findings),
            (
                f"正文中识别到 {len(data_points)} 个关键数据，其中 {provided_count} 个可在用户提供材料中找到，"
                f"{missing_count} 个未在材料中找到。"
                if data_points
                else "未发现具体数据，数据支撑偏弱。"
            ),
            _format_data_presence_lines(findings),
        )
    )

    inference_hits = _find_inference_hits(text)
    issues.append(
        _rule_issue(
            8,
            "疑似无依据推断",
            "检查是否出现“确定”“必然”“大概率上涨”“拐点已至”“配置价值凸显”等过强判断。",
            "通过" if not inference_hits else "提醒",
            "未发现明显确定性收益判断或过强投资结论。" if not inference_hits else "发现偏强判断表达，需要结合材料核实是否有依据。",
            "\n".join(inference_hits) if inference_hits else "无",
        )
    )

    product_hits = _find_product_compliance_hits(text)
    issues.append(
        _rule_issue(
            9,
            "产品合规",
            "检查是否出现其他基金公司名称 + ETF/联接/基金代码/场内/场外/A类/C类等产品词。",
            "通过" if not product_hits else "未通过",
            "未发现疑似非国泰基金产品信息。" if not product_hits else "发现疑似非国泰基金产品信息。",
            "\n".join(product_hits) if product_hits else "无",
        )
    )

    custom_status, custom_detail, custom_location = _check_article_custom(content)
    issues.append(
        _rule_issue(
            10,
            "热点文章定制项",
            "只对热点文章检查：作者是否带“国泰基金”；【相关产品】下是否只写“待填写”；【风险提示】下是否只写“待填写”。",
            custom_status,
            custom_detail,
            custom_location,
        )
    )

    return issues


def trace_content_data(content: GeneratedContentRequest) -> list[dict[str, str]]:
    article_text = "\n".join([content.summary, content.body, content.risk_notice])
    data_points = _extract_data_points(article_text)
    source_blocks = _build_source_blocks(content)
    findings: list[dict[str, str]] = []

    for token in data_points:
        location = _locate_article_token(content.body, token) or _locate_article_token(content.summary, token)
        matched = _find_token_source(token, source_blocks)
        findings.append(
            {
                "data": token,
                "article_location": location or "正文位置未定位",
                "status": matched.get("status", "未找到"),
                "source_type": matched.get("source_type", "未找到"),
                "source_location": matched.get("source_location", "未提及"),
                "source_excerpt": matched.get("source_excerpt", "无"),
            }
        )
    return findings


def generate_title_candidates(request: TitleCandidatesRequest) -> TitleCandidatesResponse:
    count = max(3, min(12, request.count))
    if settings.llm_provider.lower() in {"bailian", "dashscope"} and settings.llm_api_key:
        try:
            raw = chat_completion(_build_title_messages(request, count))
            titles = _clean_titles(_parse_titles(raw), count)
            if titles:
                return TitleCandidatesResponse(titles=titles, used_model=True)
        except BailianClientError as exc:
            return TitleCandidatesResponse(
                titles=_fallback_titles(request, count),
                used_model=False,
                issues=[str(exc)],
            )

    return TitleCandidatesResponse(titles=_fallback_titles(request, count), used_model=False)


def build_docx(content: GeneratedContentRequest) -> bytes:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    doc.add_heading(content.title or "运营内容草稿", level=1)
    if content.summary:
        doc.add_paragraph(content.summary)

    for block in _split_blocks(content.body):
        if _is_cn_heading(block):
            doc.add_heading(block, level=2)
        else:
            doc.add_paragraph(block)

    if content.risk_notice and content.risk_notice != "待填写" and content.risk_notice not in content.body:
        doc.add_heading("风险提示", level=2)
        doc.add_paragraph(content.risk_notice)

    if content.citations:
        doc.add_heading("引用与数据", level=2)
        for item in content.citations:
            doc.add_paragraph(item, style="List Bullet")

    if content.review_issues:
        doc.add_heading("质检备注", level=2)
        for issue in content.review_issues:
            doc.add_paragraph(f"{issue.level} · {issue.message}", style="List Bullet")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _issue(level: str, field: str, message: str) -> ValidationIssue:
    return ValidationIssue(level=level, field=field, message=message)


def _rule_issue(
    index: int,
    name: str,
    rule: str,
    status: str,
    detail: str,
    location: str,
) -> ValidationIssue:
    level = "info"
    if status == "提醒":
        level = "warning"
    elif status == "未通过":
        level = "error"
    return _issue(
        level,
        f"{index}. {name}",
        "\n".join(
            [
                f"结果：{status}",
                f"检测规则：{rule}",
                f"检测结果：{detail}",
                f"相关定位：{location}",
            ]
        ),
    )


def _check_structure(content: GeneratedContentRequest) -> tuple[str, str, str]:
    if content.brief_type == "热点文章":
        headings = ["【热点解读】", "【后市展望】", "【相关产品】", "【风险提示】"]
        lines = [f"- {heading}：{'已找到' if heading in content.body else '未找到'}" for heading in headings]
        missing = [heading for heading in headings if heading not in content.body]
        if missing:
            return "未通过", f"缺少 {len(missing)} 个热点文章核心模块。", "\n".join(lines)
        return "通过", "四个核心模块均已出现。", "\n".join(lines)

    keywords = ["主推方案", "机会解读", "自查清单"]
    lines = [f"- {keyword}：{'已找到' if keyword in content.body else '未找到'}" for keyword in keywords]
    missing = [keyword for keyword in keywords if keyword not in content.body]
    if missing:
        return "提醒", f"热点异动关键结构中有 {len(missing)} 项未找到。", "\n".join(lines)
    return "通过", "热点异动关键结构已出现。", "\n".join(lines)


def _find_markdown_hits(text: str) -> list[str]:
    hits: list[str] = []
    if re.search(r"(^|\n)\s*#{1,6}\s+", text):
        hits.append("- 出现 Markdown 标题符号 #")
    if text.count("**") >= 4:
        hits.append("- 出现较多 ** 加粗符号")
    if "```" in text:
        hits.append("- 出现代码块符号 ```")
    return hits


def _find_filename_hits(text: str) -> list[str]:
    pattern = re.compile(r"[\w\u4e00-\u9fff.-]*(?:WechatIMG|截屏|截图|IMG)[\w\u4e00-\u9fff.-]*|[\w\u4e00-\u9fff.-]+\.(?:jpg|jpeg|png|webp|bmp|docx|txt|md)", re.I)
    return [f"- {match.group(0)}" for match in pattern.finditer(text or "")][:20]


def _source_completeness_status(findings: list[dict[str, str]]) -> str:
    if not findings:
        return "提醒"
    missing = sum(1 for item in findings if item["status"] == "未找到")
    partial = sum(1 for item in findings if item["status"] == "找到部分来源")
    if missing:
        return "未通过"
    if partial:
        return "提醒"
    return "通过"


def _source_completeness_detail(findings: list[dict[str, str]]) -> str:
    if not findings:
        return "未识别到可溯源数据，无法检查来源完整性。"
    found = sum(1 for item in findings if item["status"] == "已找到")
    partial = sum(1 for item in findings if item["status"] == "找到部分来源")
    missing = sum(1 for item in findings if item["status"] == "未找到")
    return f"共识别 {len(findings)} 个数据，已找到来源 {found} 个，找到部分来源 {partial} 个，未找到来源 {missing} 个。"


def _data_presence_status(findings: list[dict[str, str]]) -> str:
    if not findings:
        return "提醒"
    if any(item["status"] == "未找到" for item in findings):
        return "提醒"
    return "通过"


def _format_data_presence_lines(findings: list[dict[str, str]]) -> str:
    if not findings:
        return "无"
    lines = []
    for finding in findings:
        source_result = "是，来自用户提供材料" if finding["status"] != "未找到" else "未在用户提供材料中找到"
        lines.append(f"- {finding['data']}：{source_result}")
    return "\n".join(lines)


def _format_trace_lines(findings: list[dict[str, str]]) -> str:
    if not findings:
        return "无"
    return "\n\n".join(_format_trace_message(index, finding) for index, finding in enumerate(findings, start=1))


def _check_data_caliber(findings: list[dict[str, str]]) -> tuple[str, str, str]:
    if not findings:
        return "提醒", "未识别到可检查口径的数据。", "无"

    lines: list[str] = []
    incomplete = 0
    for finding in findings:
        data = finding["data"]
        if finding["status"] == "未找到":
            incomplete += 1
            lines.append(f"- {data}：缺少来源定位，无法判断口径是否完整。")
            continue
        source_text = finding["source_excerpt"]
        has_time = bool(re.search(r"\d{4}年|截至|上半年|下半年|一季度|二季度|三季度|四季度|单日|当日|期间|同比|环比", source_text))
        has_subject = len(source_text) >= len(data) + 8
        if has_time and has_subject:
            lines.append(f"- {data}：有时间/区间或上下文主体，口径相对完整。")
        else:
            incomplete += 1
            lines.append(f"- {data}：来源片段较短或缺少时间、主体、区间等口径信息。")

    if incomplete == 0:
        return "通过", "识别到的数据均有较清晰的时间、主体或统计口径。", "\n".join(lines)
    return "提醒", f"有 {incomplete} 个数据的口径不够完整或无法判断。", "\n".join(lines)


def _find_inference_hits(text: str) -> list[str]:
    words = ["确定", "必然", "大概率上涨", "拐点已至", "配置价值凸显", "确定性", "稳赚", "必须买", "现在就买"]
    return [f"- {word}" for word in words if word in text]


def _find_product_compliance_hits(text: str) -> list[str]:
    hits: list[str] = []
    for sentence in re.split(r"[。；;\n]", text or ""):
        if "国泰基金" in sentence:
            continue
        if any(company in sentence for company in DISALLOWED_COMPANIES) and PRODUCT_WORD_RE.search(sentence):
            hits.append(f"- {sentence.strip()}")
    return hits[:20]


def _check_article_custom(content: GeneratedContentRequest) -> tuple[str, str, str]:
    if content.brief_type != "热点文章":
        return "通过", "当前不是热点文章，不适用该定制项。", "不适用"

    author_ok = "大家好，我是国泰基金" in content.body
    product_ok = bool(re.search(r"【相关产品】\s*\n+\s*待填写", content.body))
    risk_ok = bool(re.search(r"【风险提示】\s*\n+\s*待填写", content.body))
    lines = [
        f"- 作者：{'通过，已带国泰基金前缀' if author_ok else '未通过，未检测到“大家好，我是国泰基金XX。”'}",
        f"- 【相关产品】：{'通过，内容为“待填写”' if product_ok else '提醒，未检测到只保留“待填写”'}",
        f"- 【风险提示】：{'通过，内容为“待填写”' if risk_ok else '提醒，未检测到只保留“待填写”'}",
    ]
    if author_ok and product_ok and risk_ok:
        return "通过", "热点文章定制项均符合当前规则。", "\n".join(lines)
    if not author_ok:
        return "未通过", "作者前缀不符合热点文章定制要求。", "\n".join(lines)
    return "提醒", "产品或风险提示段需要人工确认是否只保留待填写。", "\n".join(lines)


DATA_POINT_RE = re.compile(
    r"(?<![A-Za-z0-9])(?:\d{4}年(?:\d{1,2}月(?:\d{1,2}日)?)?|\d{1,2}月\d{1,2}日|\d+(?:\.\d+)?\s*(?:%|％|个基点|BP|bp|万亿|万亿元|亿元|亿美元|亿|万元|美元|元|家|只|个|次|倍|点|cm|CM|年|个月|日|名|种|条)|\d+(?:\.\d+)?%)"
)


def _extract_data_points(text: str) -> list[str]:
    matches = [match.group(0).replace(" ", "") for match in DATA_POINT_RE.finditer(text or "")]
    cleaned: list[str] = []
    for value in matches:
        if value in cleaned:
            continue
        if len(value) <= 1:
            continue
        cleaned.append(value)
    return cleaned[:80]


def _build_source_blocks(content: GeneratedContentRequest) -> list[dict[str, str]]:
    blocks: list[dict[str, str]] = []
    blocks.extend(_paragraph_blocks("文字材料", "文字材料", content.material_text))
    blocks.extend(_image_blocks(content.vision_text))
    blocks.extend(_paragraph_blocks("参考文档", "参考文档", content.document_text))

    for item in content.source_data:
        value = "" if item.value is None else str(item.value)
        blocks.append(
            {
                "source_type": "Wind 数据",
                "source_location": f"{item.source or 'Wind'}｜{item.symbol}｜{item.indicator}｜{item.date}",
                "text": f"{item.symbol} {item.indicator} {item.date} {value}{item.unit}",
            }
        )
    return blocks


def _paragraph_blocks(source_type: str, label: str, text: str) -> list[dict[str, str]]:
    blocks: list[dict[str, str]] = []
    for index, paragraph in enumerate(_split_paragraphs(text), start=1):
        blocks.append(
            {
                "source_type": source_type,
                "source_location": f"{label}第{index}段",
                "text": paragraph,
            }
        )
    return blocks


def _image_blocks(text: str) -> list[dict[str, str]]:
    blocks: list[dict[str, str]] = []
    current_label = "图片材料"
    current_lines: list[str] = []
    for line in (text or "").splitlines():
        label_match = re.match(r"【图片(\d+)[:：][^】]*】", line.strip())
        if label_match:
            if current_lines:
                blocks.append(
                    {
                        "source_type": "图片材料",
                        "source_location": current_label,
                        "text": "\n".join(current_lines).strip(),
                    }
                )
            current_label = f"图片{label_match.group(1)}"
            current_lines = []
            continue
        current_lines.append(line)
    if current_lines:
        blocks.append(
            {
                "source_type": "图片材料",
                "source_location": current_label,
                "text": "\n".join(current_lines).strip(),
            }
        )
    return [block for block in blocks if block["text"]]


def _split_paragraphs(text: str) -> list[str]:
    return [paragraph.strip() for paragraph in re.split(r"\n{2,}|\r\n{2,}", text or "") if paragraph.strip()]


def _find_token_source(token: str, source_blocks: list[dict[str, str]]) -> dict[str, str]:
    normalized_token = _normalize_for_match(token)
    for block in source_blocks:
        text = block["text"]
        if token in text or normalized_token in _normalize_for_match(text):
            return {
                "status": "已找到",
                "source_type": block["source_type"],
                "source_location": block["source_location"],
                "source_excerpt": _excerpt_around(text, token),
            }
    partial = _find_partial_token_source(token, source_blocks)
    if partial:
        return partial
    return {"status": "未找到", "source_type": "未找到", "source_location": "未提及", "source_excerpt": "无"}


def _find_partial_token_source(token: str, source_blocks: list[dict[str, str]]) -> Optional[dict[str, str]]:
    numeric = re.search(r"\d+(?:\.\d+)?", token)
    if not numeric:
        return None
    number = numeric.group(0)
    for block in source_blocks:
        if number in block["text"]:
            return {
                "status": "找到部分来源",
                "source_type": block["source_type"],
                "source_location": block["source_location"],
                "source_excerpt": _excerpt_around(block["text"], number),
            }
    return None


def _normalize_for_match(value: str) -> str:
    return re.sub(r"\s+", "", value or "").replace("％", "%")


def _excerpt_around(text: str, token: str, window: int = 70) -> str:
    clean_text = re.sub(r"\s+", " ", text or "").strip()
    normalized = _normalize_for_match(clean_text)
    normalized_token = _normalize_for_match(token)
    index = normalized.find(normalized_token)
    if index < 0:
        number_match = re.search(r"\d+(?:\.\d+)?", token)
        if number_match:
            index = normalized.find(number_match.group(0))
    if index < 0:
        return clean_text[: window * 2] or "无"
    start = max(0, index - window)
    end = min(len(normalized), index + len(normalized_token) + window)
    return normalized[start:end]


def _locate_article_token(body: str, token: str) -> str:
    if not body or token not in body:
        normalized_body = _normalize_for_match(body)
        if _normalize_for_match(token) not in normalized_body:
            return ""

    current_section = "开场/摘要"
    section_counts: dict[str, int] = {}
    paragraphs = _split_paragraphs(body)
    normalized_token = _normalize_for_match(token)
    for paragraph in paragraphs:
        heading_match = re.fullmatch(r"【[^】]{2,12}】", paragraph.strip())
        if heading_match:
            current_section = paragraph.strip()
            section_counts.setdefault(current_section, 0)
            continue
        section_counts[current_section] = section_counts.get(current_section, 0) + 1
        if token in paragraph or normalized_token in _normalize_for_match(paragraph):
            return f"{current_section}第{section_counts[current_section]}段"
    return "正文位置未定位"


def _format_trace_message(index: int, finding: dict[str, str]) -> str:
    return (
        f"规则{index}｜数据溯源｜数据：{finding['data']}\n"
        f"正文位置：{finding['article_location']}\n"
        f"来源结果：{finding['status']}\n"
        f"来源类型：{finding['source_type']}\n"
        f"来源定位：{finding['source_location']}\n"
        f"来源片段：{finding['source_excerpt']}"
    )


def _score_from_trace(findings: list[dict[str, str]]) -> int:
    if not findings:
        return 70
    found = sum(1 for item in findings if item["status"] == "已找到")
    partial = sum(1 for item in findings if item["status"] == "找到部分来源")
    missing = sum(1 for item in findings if item["status"] == "未找到")
    score = int(((found * 1.0 + partial * 0.55) / len(findings)) * 100)
    if missing:
        score = max(0, score - min(18, missing * 3))
    return max(0, min(100, score))


def _score_from_issues(issues: list[ValidationIssue]) -> int:
    score = 100
    for issue in issues:
        if issue.level == "error":
            score -= 18
        elif issue.level == "warning":
            score -= 7
    return max(0, min(100, score))


def _dedupe_issues(issues: list[ValidationIssue]) -> list[ValidationIssue]:
    seen: set[tuple[str, str, str]] = set()
    result: list[ValidationIssue] = []
    for issue in issues:
        key = (issue.level, issue.field, issue.message)
        if key in seen:
            continue
        seen.add(key)
        result.append(issue)
    return result


def _build_title_messages(request: TitleCandidatesRequest, count: int) -> list[dict[str, str]]:
    return [
        {
            "role": "system",
            "content": (
                "你是基金运营标题助手。只能基于用户提供的主题、正文和材料生成标题；"
                "不得出现收益承诺、确定性判断、非国泰基金产品名称或其他基金公司产品。"
                "标题要有信息密度，问句可以使用，但不要每条都写“后市怎么看”。"
                "只输出合法 JSON：{\"titles\":[\"标题1\"]}。"
            ),
        },
        {
            "role": "user",
            "content": json.dumps(
                {
                    "brief_type": request.brief_type,
                    "count": count,
                    "topic": request.topic,
                    "title_hint": request.title_hint,
                    "title_prompt": request.title_prompt,
                    "summary": request.summary,
                    "body": request.body[:6000],
                    "material_text": request.material_text[:4000],
                },
                ensure_ascii=False,
            ),
        },
    ]


def _parse_titles(raw: str) -> list[str]:
    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", raw, re.S)
        if match:
            try:
                parsed = json.loads(match.group(0))
            except json.JSONDecodeError:
                parsed = raw
        else:
            parsed = raw
    if isinstance(parsed, dict) and isinstance(parsed.get("titles"), list):
        return [str(item) for item in parsed["titles"]]
    if isinstance(parsed, list):
        return [str(item) for item in parsed]
    return [line.strip(" -0123456789.、") for line in str(parsed).splitlines()]


def _clean_titles(titles: list[str], count: int) -> list[str]:
    cleaned: list[str] = []
    for title in titles:
        value = re.sub(r"^[\"'“”]+|[\"'“”]+$", "", title.strip())
        value = re.sub(r"#{1,6}\s*", "", value)
        if not value or value in cleaned:
            continue
        if any(word in value for word in FORBIDDEN_TITLE_WORDS):
            continue
        if any(company in value for company in DISALLOWED_COMPANIES) and PRODUCT_WORD_RE.search(value):
            continue
        cleaned.append(value)
        if len(cleaned) >= count:
            break
    return cleaned


def _fallback_titles(request: TitleCandidatesRequest, count: int) -> list[str]:
    topic = (request.topic or request.title_hint or "热点主题").strip()[:18]
    if request.brief_type == "热点文章":
        candidates = [
            f"{topic}关注度升温，关键变量如何演绎？",
            f"{topic}迎来新变化，景气修复能否延续？",
            f"{topic}多重信号共振，市场关注点在哪？",
            f"{topic}再成焦点，产业逻辑出现哪些变化？",
            f"{topic}热度抬升，后续催化如何观察？",
            f"{topic}行情升温，核心支撑来自哪里？",
            f"{topic}站上风口，政策与产业如何共振？",
            f"{topic}迎来窗口期，关注哪些验证信号？",
        ]
    else:
        candidates = [
            f"{topic}出现异动，关注核心催化变化",
            f"{topic}热度提升，相关逻辑值得跟踪",
            f"{topic}阶段性走强，重点关注数据验证",
            f"{topic}市场关注升温，短期催化增强",
            f"{topic}迎来资金关注，后续观察持续性",
            f"{topic}波动放大，运营素材可重点跟进",
        ]
    return candidates[:count]


def _split_blocks(text: str) -> list[str]:
    return [block.strip() for block in re.split(r"\n{2,}", text.strip()) if block.strip()]


def _is_cn_heading(block: str) -> bool:
    return bool(re.fullmatch(r"【[^】]{2,12}】", block.strip()))
